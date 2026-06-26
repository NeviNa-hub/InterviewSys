import hashlib
import os
import re
from typing import Iterable

from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from pypdf import PdfReader

from utils.logger_handler import logger


def ensure_dir(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path


def slugify_filename(filename: str) -> str:
    name = os.path.basename(filename or "file")
    stem, ext = os.path.splitext(name)
    safe_stem = re.sub(r"[^a-zA-Z0-9_\-\u4e00-\u9fff]+", "_", stem).strip("_") or "file"
    safe_ext = re.sub(r"[^a-zA-Z0-9.]+", "", ext.lower())
    return f"{safe_stem}{safe_ext}"


def save_uploaded_file(uploaded_file, target_dir: str) -> str:
    ensure_dir(target_dir)
    file_name = slugify_filename(getattr(uploaded_file, "name", "upload.bin"))
    save_path = os.path.join(target_dir, file_name)
    with open(save_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return save_path


def save_binary_file(filename: str, data: bytes, target_dir: str) -> str:
    ensure_dir(target_dir)
    file_name = slugify_filename(filename or "upload.bin")
    save_path = os.path.join(target_dir, file_name)
    with open(save_path, "wb") as f:
        f.write(data)
    return save_path


def get_file_md5_hex(filepath: str):
    if not os.path.exists(filepath):
        logger.error(f"[md5] 文件不存在: {filepath}")
        return None

    if not os.path.isfile(filepath):
        logger.error(f"[md5] 路径不是文件: {filepath}")
        return None

    md5_obj = hashlib.md5()
    chunk_size = 4096
    try:
        with open(filepath, "rb") as f:
            while True:
                chunk = f.read(chunk_size)
                if not chunk:
                    break
                md5_obj.update(chunk)
        return md5_obj.hexdigest()
    except Exception as e:
        logger.error(f"[md5] 计算文件 {filepath} 失败: {str(e)}")
        return None


def listdir_with_allowed_type(path: str, allowed_types: tuple[str]):
    files = []
    if not os.path.isdir(path):
        logger.error(f"[listdir_with_allowed_type] {path} 不是文件夹")
        return tuple(files)

    normalized_types = tuple(ext.lower().lstrip(".") for ext in allowed_types)
    for root, _, filenames in os.walk(path):
        for file_name in filenames:
            ext = os.path.splitext(file_name)[1].lower().lstrip(".")
            if ext in normalized_types:
                files.append(os.path.join(root, file_name))
    return tuple(files)


def pdf_loader(filepath: str, passwd=None) -> list[Document]:
    return PyPDFLoader(filepath, passwd).load()


def txt_loader(filepath: str) -> list[Document]:
    return TextLoader(filepath, encoding="utf-8").load()


def md_loader(filepath: str) -> list[Document]:
    return TextLoader(filepath, encoding="utf-8").load()


def docx_loader(filepath: str) -> list[Document]:
    doc = DocxDocument(filepath)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
    return [Document(page_content=text, metadata={"source": filepath})]


def extract_text_from_file(filepath: str) -> str:
    suffix = os.path.splitext(filepath)[1].lower()

    if suffix in {".txt", ".md"}:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read().strip()

    if suffix == ".pdf":
        reader = PdfReader(filepath)
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()

    if suffix == ".docx":
        doc = DocxDocument(filepath)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]).strip()

    raise ValueError(f"暂不支持解析该文件类型: {suffix}")


def unique_non_empty_texts(items: Iterable[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        text = str(item or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        result.append(text)
    return result
